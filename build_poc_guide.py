"""Build the Signature Sprint POC Guide as a branded Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1E, 0x2A, 0x3D)   # Hyreo "hy" dark
CYAN        = RGBColor(0x3B, 0xD4, 0xE7)   # Hyreo "reo" accent
INK         = RGBColor(0x1F, 0x2A, 0x37)   # body text
MUTED       = RGBColor(0x64, 0x74, 0x8B)   # captions / meta
RULE        = RGBColor(0xE5, 0xEA, 0xF0)   # hairline
BG_SOFT     = "F5F9FC"                     # soft callout bg (hex string for shading)
BG_ACCENT   = "E8FAFD"                     # accent callout bg

BODY_FONT    = "Calibri"
DISPLAY_FONT = "Calibri"

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, *, top=None, bottom=None, left=None, right=None, color="E5EAF0", sz="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), val if val else "nil")
        if val:
            el.set(qn("w:sz"), sz)
            el.set(qn("w:color"), color)
        borders.append(el)
    tc_pr.append(borders)

def add_horizontal_rule(doc, color=RULE, space_before=4, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p_pr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2]))
    pbdr.append(bottom)
    p_pr.append(pbdr)

def styled_run(paragraph, text, *, font=BODY_FONT, size=11, bold=False, italic=False, color=INK):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    return r

def add_hyreo_wordmark(paragraph, size=28):
    """Render the hyreo wordmark: 'hy' in navy + 'reo' in cyan, lowercase."""
    hy = paragraph.add_run("hy")
    hy.font.name = DISPLAY_FONT
    hy.font.size = Pt(size)
    hy.bold = True
    hy.font.color.rgb = NAVY
    reo = paragraph.add_run("reo")
    reo.font.name = DISPLAY_FONT
    reo.font.size = Pt(size)
    reo.bold = True
    reo.font.color.rgb = CYAN

def add_heading(doc, text, *, level=1, space_before=18, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    sizes = {1: 20, 2: 15, 3: 12}
    colors = {1: NAVY, 2: NAVY, 3: NAVY}
    styled_run(p, text, size=sizes.get(level, 12), bold=True, color=colors.get(level, INK))
    return p

def add_body(doc, text, *, italic=False, size=11, space_after=6, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.35
    styled_run(p, text, italic=italic, size=size, color=color)
    return p

def add_bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    if bold_lead:
        styled_run(p, bold_lead, bold=True, color=NAVY)
        styled_run(p, text, color=INK)
    else:
        styled_run(p, text, color=INK)
    return p

def add_callout(doc, title, body, *, bg=BG_SOFT, accent=CYAN):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_borders(
        cell,
        top="single", bottom="single", right="single",
        left="single", color="{:02X}{:02X}{:02X}".format(accent[0], accent[1], accent[2]),
        sz="18",
    )
    cell.text = ""
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    styled_run(p_title, title, bold=True, size=11, color=NAVY)
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.3
    styled_run(p_body, body, size=10.5, color=INK)
    # Spacer after
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_info_table(doc, rows, *, widths=None, header=True):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.autofit = False
    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = w
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            is_header = header and r_idx == 0
            styled_run(
                p, str(val),
                size=10.5,
                bold=is_header,
                color=(RGBColor(0xFF, 0xFF, 0xFF) if is_header else INK),
            )
            if is_header:
                set_cell_bg(cell, "1E2A3D")
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            set_cell_borders(
                cell,
                top="single", bottom="single", left="single", right="single",
                color="E5EAF0", sz="4",
            )
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return tbl

# ── Build document ───────────────────────────────────────────────────────────
doc = Document()

# Page & default font
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(11)
style.font.color.rgb = INK

# ── Cover page ───────────────────────────────────────────────────────────────
# Logo (wordmark)
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_p.paragraph_format.space_before = Pt(24)
logo_p.paragraph_format.space_after = Pt(6)
add_hyreo_wordmark(logo_p, size=34)

tag_p = doc.add_paragraph()
tag_p.paragraph_format.space_after = Pt(80)
styled_run(tag_p, "Hyreo Labs  ·  Signature Sprint", size=10.5, color=MUTED)

# Eyebrow
eyebrow = doc.add_paragraph()
eyebrow.paragraph_format.space_after = Pt(6)
styled_run(eyebrow, "PROOF OF CONCEPT   ·   GUIDE", size=10, bold=True, color=CYAN)

# Title
title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(4)
styled_run(title, "Signature Sprint", size=34, bold=True, color=NAVY)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(18)
styled_run(
    subtitle,
    "A behavioral simulation that measures Drive for Results in under five minutes.",
    size=13, italic=True, color=MUTED,
)

add_horizontal_rule(doc, space_before=6, space_after=14)

# Meta block
meta_rows = [
    ("Product",    "Signature Sprint — Behavioral Simulation for Drive for Results"),
    ("Version",    "v2.7  ·  April 2026"),
    ("Status",     "POC  ·  Stable"),
    ("Prepared by","Hyreo Labs"),
    ("Audience",   "Evaluators, Hiring Managers, L&D Leads, POC Stakeholders"),
]
meta_tbl = doc.add_table(rows=len(meta_rows), cols=2)
meta_tbl.autofit = False
for i, (k, v) in enumerate(meta_rows):
    kc = meta_tbl.cell(i, 0); vc = meta_tbl.cell(i, 1)
    kc.width = Inches(1.7); vc.width = Inches(4.6)
    kc.text = ""; vc.text = ""
    pk = kc.paragraphs[0]; pv = vc.paragraphs[0]
    pk.paragraph_format.space_after = Pt(2)
    pv.paragraph_format.space_after = Pt(2)
    styled_run(pk, k.upper(), size=9.5, bold=True, color=MUTED)
    styled_run(pv, v, size=11, color=INK)
    set_cell_borders(kc, bottom="single", color="E5EAF0", sz="4")
    set_cell_borders(vc, bottom="single", color="E5EAF0", sz="4")

# Footer-ish note on cover
note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(120)
styled_run(note, "Confidential  ·  For evaluation partners of Hyreo Labs", size=9, italic=True, color=MUTED)

doc.add_page_break()

# ── Contents ─────────────────────────────────────────────────────────────────
add_heading(doc, "Contents", level=1, space_before=0)
toc_items = [
    "1.  Executive Summary",
    "2.  What the POC Proves",
    "3.  Product at a Glance",
    "4.  How the Assessment Works",
    "5.  The Candidate Experience",
    "6.  The Evaluator Experience",
    "7.  Scoring, Validation & Flags",
    "8.  Running the POC",
    "9.  Success Criteria",
    "10.  FAQ & Support",
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    styled_run(p, item, size=11, color=INK)

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────────
add_heading(doc, "1.  Executive Summary", level=1, space_before=0)
add_body(
    doc,
    "Signature Sprint is a five-minute, game-like behavioral simulation that measures "
    "the Drive for Results competency with anti-faking validation baked into the scoring engine. "
    "This proof of concept exists to demonstrate, with real candidates and real evaluators, "
    "that a short immersive experience can produce a cleaner behavioral signal than traditional "
    "SJTs or self-report inventories — and that evaluators can act on that signal in a two-minute read.",
)
add_body(
    doc,
    "The POC ships as a fully self-contained web app. Candidates complete four levels in the "
    "browser; evaluators review scored reports through a dedicated manager portal. No backend is "
    "required for evaluation; data persists locally in the browser for 24 hours.",
)

add_callout(
    doc,
    "What success looks like for this POC",
    "Evaluators can scan a candidate report in under two minutes, form an interview hypothesis, "
    "and identify at least one probe area. Candidates complete the sprint in a single sitting "
    "without instructional support.",
    bg=BG_ACCENT,
)

# ── 2. What the POC Proves ───────────────────────────────────────────────────
add_heading(doc, "2.  What the POC Proves", level=1)
add_body(doc, "Three hypotheses guide this proof of concept:")
add_bullet(doc, "Candidates engage through completion when the assessment feels like a game, not a form.", bold_lead="Engagement.  ")
add_bullet(doc, "Mixing tactical (real-time) and reflective (self-concept) measurement surfaces the gap between what candidates say and what they do.", bold_lead="Signal quality.  ")
add_bullet(doc, "A clean, opinionated report lets an evaluator form an interview hypothesis in under two minutes.", bold_lead="Evaluator efficiency.  ")

# ── 3. Product at a Glance ───────────────────────────────────────────────────
add_heading(doc, "3.  Product at a Glance", level=1)

glance_rows = [
    ("Attribute",           "Value"),
    ("Competency measured", "Drive for Results"),
    ("Dimensions",          "Ownership · Action Orientation · Perseverance"),
    ("Duration",            "≈ 5 minutes end-to-end"),
    ("Levels",              "4 (Execution Lab, Character Mirror, Priority Lens, Core Narrative)"),
    ("Themes",              "Adventure (warm) · Mission Control (dark)"),
    ("Roles",               "Candidate · Evaluator / Manager"),
    ("Persistence",         "Browser localStorage · 24-hour TTL"),
    ("Deployment",          "Static site — run locally, serve via HTTP, or deploy to Vercel"),
]
add_info_table(doc, glance_rows, widths=[Inches(2.0), Inches(4.3)])

add_body(doc, "")  # spacer

add_callout(
    doc,
    "Brand note",
    "The product was formerly called Catalyst Protocol. Internal code identifiers (window.CatalystCore, "
    "CatalystAdventure, CatalystMission) retain the original names for compatibility; all user-facing "
    "surfaces are now Signature Sprint.",
)

# ── 4. How the Assessment Works ──────────────────────────────────────────────
add_heading(doc, "4.  How the Assessment Works", level=1)
add_body(
    doc,
    "Drive for Results is decomposed into three weighted dimensions. Each candidate response is "
    "mapped to a Behaviorally Anchored Rating Scale (BARS) level, then rolled up into a final "
    "0–4 score.",
)

dim_rows = [
    ("Dimension",             "Weight", "What it captures"),
    ("Ownership",             "0.40",   "Does the candidate close the loop themselves?"),
    ("Action Orientation",    "0.35",   "Do they convert ambiguity into forward motion?"),
    ("Perseverance",          "0.25",   "Do they stay with hard problems past the first ceiling?"),
]
add_info_table(doc, dim_rows, widths=[Inches(1.8), Inches(1.0), Inches(3.5)])

add_body(doc, "")

add_heading(doc, "The four levels", level=2)
level_rows = [
    ("Level",              "Name",               "Modality",                  "Scenarios", "Timer"),
    ("L1",                 "The Execution Lab",  "Pick one of four moves",    "6",         "20s each"),
    ("L2",                 "The Character Mirror","More / Less Like Me",      "9",         "15s per card"),
    ("L3",                 "The Priority Lens",  "Rank 1st / 2nd / 3rd",      "3",         "25s each"),
    ("L4",                 "The Core Narrative", "Pick one of three reflections","3",      "20s each"),
]
add_info_table(doc, level_rows, widths=[Inches(0.6), Inches(1.7), Inches(2.1), Inches(0.9), Inches(1.0)])

add_body(doc, "")

add_heading(doc, "BARS bands", level=2)
bars_rows = [
    ("Band", "Label"),
    ("1",    "Ineffective"),
    ("2",    "Developing"),
    ("3",    "Effective"),
    ("4",    "Advanced"),
]
add_info_table(doc, bars_rows, widths=[Inches(0.8), Inches(2.4)])

# ── 5. Candidate Experience ──────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "5.  The Candidate Experience", level=1, space_before=0)
add_body(
    doc,
    "The candidate journey is designed around a single principle: produce an authentic behavioral "
    "signal by making the experience feel unlike an assessment. Candidates never see their score, "
    "BARS mapping, or validation flags — transparency here would invite gaming.",
)

add_heading(doc, "Journey", level=2)
add_bullet(doc, "Choose a theme (Adventure or Mission Control) from the hub.", bold_lead="Entry.  ")
add_bullet(doc, "Read a short welcome (purpose, ground rules, name entry), then a 3-2-1-GO countdown.", bold_lead="Intro.  ")
add_bullet(doc, "Each level opens with a briefing card — headline, \"The Script\", \"Your Task\", three guidelines, CTA.", bold_lead="Level briefings.  ")
add_bullet(doc, "Confetti and the level title between levels; no \"badge earned\" language.", bold_lead="Transitions.  ")
add_bullet(doc, "Four completion icons, a \"Your Signature\" paragraph, three reinforcing insights, and an Export Report button.", bold_lead="Final screen.  ")

add_heading(doc, "UX principles at work", level=2)
ux_rows = [
    ("Principle",                 "How it's expressed"),
    ("Instant content load",      "No typewriter effects; prompts render atomically"),
    ("No correct/incorrect labels","System-style feedback: Move logged · Response locked in · Reflection saved"),
    ("Smooth over snappy",        "Weighted transitions, staggered options, spring bounce on pick"),
    ("One decision at a time",    "Single prompt in a prominent chip above each level"),
    ("Timer as urgency, not punishment","TIME UP overlay auto-logs the most cautious option and proceeds"),
    ("Celebrate every level",     "Confetti + level title between levels"),
]
add_info_table(doc, ux_rows, widths=[Inches(2.0), Inches(4.3)])

# ── 6. Evaluator Experience ──────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "6.  The Evaluator Experience", level=1, space_before=0)
add_body(
    doc,
    "Evaluators work from a dedicated manager portal accessible from either theme. Both themes "
    "render the same report — clean Inter typography, print-optimized for interview prep.",
)

add_heading(doc, "The candidate list", level=2)
add_bullet(doc, "Sortable table: Name · Date · Score · Level band.", bold_lead="Columns.  ")
add_bullet(doc, "Total candidates · average score · count at Level 4 · flagged count.", bold_lead="KPI row.  ")
add_bullet(doc, "Search by name; sort by date (default), score, or name.", bold_lead="Controls.  ")
add_bullet(doc, "A ⚠ icon next to a name signals one or more validation flags on that run.", bold_lead="Flag indicator.  ")

add_heading(doc, "Individual report — what you see", level=2)
add_bullet(doc, "Avatar, name, date, final score, and L1–L4 level badge.", bold_lead="Header.  ")
add_bullet(doc, "Amber cards, shown only when triggered. Conversation starters, not verdicts.", bold_lead="Validation flags.  ")
add_bullet(doc, "Three animated horizontal bars — Ownership · Action Orientation · Perseverance.", bold_lead="Dimension breakdown.  ")
add_bullet(doc, "Average response time, fast-response count, consistency check, temperament pattern (Snap / Balanced / Deliberate).", bold_lead="Validation signals.  ")
add_bullet(doc, "Four per-level tiles; tiles turn amber when > 0. Total-count chip in the card header.", bold_lead="Skipped questions.  ")
add_bullet(doc, "Highest dimension with a note; weakest dimension with a concrete development tip.", bold_lead="Strengths & growth.  ")
add_bullet(doc, "One-paragraph interpretation tied to the candidate's band.", bold_lead="Behavioral summary.  ")
add_bullet(doc, "Opens the browser's print dialog; layout is print-optimized.", bold_lead="Export.  ")

add_callout(
    doc,
    "Two-minute read",
    "1. Look at the score first — L1 red, L2 amber, L3 blue, L4 green.  "
    "2. Scan for flags — amber cards = probe areas.  "
    "3. Check the dimension bars — the weakest is the growth edge.  "
    "4. Skim the behavioral summary for the one-paragraph verdict.  "
    "5. Export and bring the PDF into interview prep.",
    bg=BG_ACCENT,
)

# ── 7. Scoring, Validation & Flags ───────────────────────────────────────────
add_heading(doc, "7.  Scoring, Validation & Flags", level=1)
add_body(doc, "The final score is the weighted sum of the three dimensions:")

formula_p = doc.add_paragraph()
formula_p.paragraph_format.left_indent = Inches(0.3)
formula_p.paragraph_format.space_before = Pt(4)
formula_p.paragraph_format.space_after = Pt(8)
styled_run(
    formula_p,
    "FinalScore = (Ownership × 0.40) + (Action Orientation × 0.35) + (Perseverance × 0.25)",
    font="Consolas", size=10.5, color=NAVY,
)

add_body(
    doc,
    "Validation signals sit alongside the score to help evaluators read context, not to disqualify. "
    "The report surfaces response-time patterns, consistency between tactical and self-concept "
    "levels, and skipped-question counts per level.",
)

add_callout(
    doc,
    "Flags are conversation starters",
    "A flagged candidate is not disqualified. Each flag points to a specific behavior worth probing "
    "in the interview. See the Report Reference for each flag's trigger logic and suggested follow-up "
    "questions.",
)

# ── 8. Running the POC ───────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "8.  Running the POC", level=1, space_before=0)

add_heading(doc, "Quickest — open directly", level=2)
add_body(doc, "Double-click index.html. Works in any modern browser (Chrome, Firefox, Edge, Safari).")
add_callout(
    doc,
    "Caveat",
    "On a file:// origin, Chrome isolates localStorage per file, so runs completed in Adventure "
    "won't show up in Mission's Manager view and vice versa. For consistent cross-theme behavior, "
    "serve over HTTP.",
)

add_heading(doc, "Serve over HTTP", level=2)
code_p = doc.add_paragraph()
code_p.paragraph_format.left_indent = Inches(0.3)
code_p.paragraph_format.space_after = Pt(2)
styled_run(code_p, "cd d:\\BAT", font="Consolas", size=10.5, color=NAVY)
code_p2 = doc.add_paragraph()
code_p2.paragraph_format.left_indent = Inches(0.3)
code_p2.paragraph_format.space_after = Pt(2)
styled_run(code_p2, "python -m http.server 8000", font="Consolas", size=10.5, color=NAVY)
code_p3 = doc.add_paragraph()
code_p3.paragraph_format.left_indent = Inches(0.3)
code_p3.paragraph_format.space_after = Pt(8)
styled_run(code_p3, "# or: npx serve .", font="Consolas", size=10.5, color=MUTED)
add_body(doc, "Then open http://localhost:8000/ — both themes share the same origin and localStorage.")

add_heading(doc, "Deploy to Vercel", level=2)
add_bullet(doc, "From the project directory, run `vercel` and then `vercel --prod`.", bold_lead="Vercel CLI.  ")
add_bullet(doc, "Push to a GitHub repo and import it from vercel.com/new.", bold_lead="GitHub integration.  ")
add_bullet(doc, "Zip the folder's contents (not the folder itself) and drop onto vercel.com/new.", bold_lead="Drag-and-drop.  ")
add_body(doc, "No build step is required — the app runs on Babel Standalone in the browser.", italic=True, color=MUTED, size=10.5)

# ── 9. Success Criteria ──────────────────────────────────────────────────────
add_heading(doc, "9.  Success Criteria", level=1)
add_body(
    doc,
    "The POC is considered successful when the following signals are present across the pilot cohort:",
)
sc_rows = [
    ("Signal",                              "Target"),
    ("Candidate completion rate",           "≥ 90% of started sprints reach the final screen"),
    ("Candidate time-on-task",              "Median under 6 minutes; 95th percentile under 9"),
    ("Evaluator time-to-hypothesis",        "Under 2 minutes per report on first reading"),
    ("Evaluator flag action rate",          "Each surfaced flag maps to an interview probe"),
    ("Signal divergence (L1 vs L2)",        "Detectable gap between tactical and self-concept responses"),
]
add_info_table(doc, sc_rows, widths=[Inches(2.6), Inches(3.7)])

# ── 10. FAQ & Support ────────────────────────────────────────────────────────
add_heading(doc, "10.  FAQ & Support", level=1)

faqs = [
    ("I clicked the wrong option — can I go back?",
     "No. The Sprint is one-way to preserve honest first-instinct signal. Start a new session from the home hub to replay."),
    ("What if the timer runs out?",
     "A brief TIME UP · SCENARIO SKIPPED banner appears and the system auto-logs the most cautious option. Time-outs are a signal, not a failure."),
    ("Who sees my answers?",
     "Only the evaluator, via the Manager Portal. At POC scale, data stays in the browser for 24 hours."),
    ("Why don't candidates see their score?",
     "Candidates who see their scores start gaming future picks. Hiding the score preserves the authentic behavioral signal."),
    ("Alex or Avantika — does it matter which is picked in L2?",
     "No. Both characters have the same 9 cards and scenarios. The choice is purely cosmetic."),
    ("The Manager dashboard shows 5 unfamiliar names.",
     "Those are seed mock candidates (Ananya, Rohan, Priya, Kabir, Meera) for demos. Real completed runs appear above them."),
    ("A sprint from yesterday is missing from the Manager list.",
     "The 24-hour TTL rolled that run off. Complete a fresh sprint — it will persist for another 24 hours."),
    ("Can we add more scenarios or edit questions?",
     "Yes. Edit the L1, L2, L3, or L4 arrays in catalyst-shared.js. Both themes consume the same source; changes flow to both on reload."),
]
for q, a in faqs:
    qp = doc.add_paragraph()
    qp.paragraph_format.space_before = Pt(6)
    qp.paragraph_format.space_after = Pt(2)
    styled_run(qp, "Q.  ", bold=True, color=CYAN, size=11)
    styled_run(qp, q, bold=True, color=NAVY, size=11)
    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(4)
    ap.paragraph_format.left_indent = Inches(0.25)
    styled_run(ap, a, size=10.5, color=INK)

add_body(doc, "")
add_horizontal_rule(doc, space_before=10, space_after=6)

# ── Document footer block ────────────────────────────────────────────────────
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
foot.paragraph_format.space_after = Pt(2)
add_hyreo_wordmark(foot, size=14)
foot2 = doc.add_paragraph()
styled_run(foot2, "Hyreo Labs  ·  Signature Sprint POC Guide  ·  v2.7  ·  April 2026", size=9.5, italic=True, color=MUTED)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r"d:\BAT\Signature-Sprint-POC-Guide.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
